"""Render the project's Markdown deliverables as clean .docx for sharing.

The repo keeps Markdown because that's what GitHub renders. This produces a Word
copy of the same content, in one consistent house style.

    python tools/make_docx.py                    # all deliverables
    python tools/make_docx.py BUGS.md            # just one
    python tools/make_docx.py LOOM1.md --compact # squeeze onto a single page

Page counts can only be measured by rendering; estimating from line counts got the
narration sheets wrong. To check:

    soffice --headless --convert-to pdf --outdir /tmp *.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DEFAULTS = ["README.md", "ARCHITECTURE.md", "BUGS.md", "SUBMISSION.md"]

BODY_FONT = "Calibri"
MONO_FONT = "Menlo"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
LINK = RGBColor(0x1A, 0x5F, 0xB4)
RULE_GREY = "D8D8D8"
CODE_BG = "F4F4F4"
HEAD_BG = "EFEFEF"

TEXT_WIDTH = Inches(6.9)  # letter minus 0.8in margins each side

# `code` | **bold** | *italic* | [text](url). Order matters: code first so backticked
# markup stays literal, and bold before italic so ** isn't read as a single *.
INLINE = re.compile(
    r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*\s][^*]*\*|\[[^\]]+\]\([^)]+\))"
)
BULLET = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBERED = re.compile(r"^(\s*)\d+\.\s+(.*)$")
HR = {"---", "***", "___"}


# --------------------------------------------------------------------------- #
# Low-level docx helpers (python-docx has no API for any of these)
# --------------------------------------------------------------------------- #


def _shade(element, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def shade_cell(cell, fill: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    _shade(paragraph._p.get_or_add_pPr(), fill)


def add_border(paragraph, edge: str, size: int = 6, color: str = RULE_GREY) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = pPr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        pPr.append(borders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "6")
    el.set(qn("w:color"), color)
    borders.append(el)


def add_hyperlink(paragraph, url: str, label: str, size: Pt) -> None:
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag, attrs in (
        ("w:color", {"w:val": "1A5FB4"}),
        ("w:u", {"w:val": "single"}),
        ("w:rFonts", {"w:ascii": BODY_FONT, "w:hAnsi": BODY_FONT}),
        ("w:sz", {"w:val": str(int(size.pt * 2))}),
    ):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        rPr.append(el)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = label
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


# --------------------------------------------------------------------------- #
# Inline runs
# --------------------------------------------------------------------------- #


def add_inline(
    paragraph, text: str, size: Pt, bold: bool = False, italic: bool = False
) -> None:
    """Emit runs for one line of inline Markdown.

    Recurses, because the markup nests: a bold beat label containing a filename in
    backticks used to render the backticks literally, and *emphasis* was passed
    through untouched.
    """
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(size.pt - 1.0)
            run.bold = bold
        elif piece.startswith("**") and piece.endswith("**"):
            add_inline(paragraph, piece[2:-2], size, bold=True, italic=italic)
        elif piece.startswith("*") and piece.endswith("*"):
            add_inline(paragraph, piece[1:-1], size, bold=bold, italic=True)
        elif link := re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", piece):
            label, url = link.groups()
            if "://" in url:
                # A hyperlink is a single run, so drop any markup in the label.
                add_hyperlink(paragraph, url, label.replace("`", ""), size)
            else:  # in-repo path: show it, don't pretend it's clickable
                start = len(paragraph.runs)
                add_inline(paragraph, label, size, bold=bold, italic=italic)
                for run in paragraph.runs[start:]:
                    run.font.color.rgb = LINK
        else:
            run = paragraph.add_run(piece)
            run.bold = bold
            run.italic = italic


# --------------------------------------------------------------------------- #
# Block elements
# --------------------------------------------------------------------------- #


def add_prose(doc, text: str, size: Pt) -> None:
    add_inline(doc.add_paragraph(), text, size)


def add_bullet(doc, text: str, size: Pt, level: int) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    add_inline(p, text, size)


def add_numbered(doc, text: str, size: Pt) -> None:
    add_inline(doc.add_paragraph(style="List Number"), text, size)


def add_quote(doc, text: str, size: Pt) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.28)
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(4)
    add_border(p, "left", size=12, color="C8C8C8")
    add_inline(p, text, size, italic=True)
    for run in p.runs:
        if not run.bold:
            run.font.color.rgb = MUTED


def add_code(doc, lines: list[str], size: Pt) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.14)
    fmt.space_before = Pt(5)
    fmt.space_after = Pt(7)
    shade_paragraph(p, CODE_BG)
    for n, line in enumerate(lines):
        if n:
            p.add_run().add_break()  # real line breaks, not \n inside a run
        run = p.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(size.pt - 1.5)
        run.font.color.rgb = INK


def add_rule(doc) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    add_border(p, "bottom")


def add_table(doc, rows: list[list[str]], size: Pt) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Column widths proportional to content, so a "#" column stays narrow.
    weights = [max(len(r[c]) for r in rows) for c in range(len(rows[0]))]
    floor = 0.55 / (TEXT_WIDTH.inches)
    share = [max(w / sum(weights), floor) for w in weights]
    share = [s / sum(share) for s in share]

    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            cell.width = Inches(TEXT_WIDTH.inches * share[c])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text, Pt(size.pt - 0.5))
            if r == 0:
                shade_cell(cell, HEAD_BG)
                for run in p.runs:
                    run.bold = True
    # Repeat the header if the table breaks across pages.
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# --------------------------------------------------------------------------- #
# Parsing
# --------------------------------------------------------------------------- #


def is_divider(line: str) -> bool:
    s = line.strip()
    return "-" in s and bool(re.fullmatch(r"\|?[\s:|-]+\|?", s))


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def starts_block(line: str) -> bool:
    """Does this line begin a new block, rather than continue the current one?"""
    s = line.strip()
    return (
        not s
        or s.startswith(("#", ">", "```", "|"))
        or s in HR
        or bool(BULLET.match(line))
        or bool(NUMBERED.match(line))
    )


def setup_styles(doc: Document, compact: bool) -> Pt:
    size = Pt(9.5 if compact else 10.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5 if compact else 0.75)
        section.left_margin = section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = size
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2 if compact else 6)
    normal.paragraph_format.line_spacing = 1.0 if compact else 1.12

    heads = (
        [("Heading 1", 14, 0, 5), ("Heading 2", 11.5, 8, 3), ("Heading 3", 10, 6, 2)]
        if compact
        else [("Heading 1", 18, 0, 9), ("Heading 2", 13.5, 16, 6), ("Heading 3", 11.5, 12, 4)]
    )
    for name, pt, before, after in heads:
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(pt)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = INK  # Word defaults these to blue
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        try:
            fmt = doc.styles[name].paragraph_format
        except KeyError:
            continue
        fmt.space_after = Pt(1 if compact else 3)
        fmt.line_spacing = 1.0 if compact else 1.12
    return size


def convert(md_path: Path, compact: bool = False) -> Path:
    doc = Document()
    size = setup_styles(doc, compact)

    lines = md_path.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if s.startswith("```"):
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code(doc, block, size)
            i += 1

        elif s.startswith("|") and i + 1 < len(lines) and is_divider(lines[i + 1]):
            rows = [split_row(s)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows, size)

        elif s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            doc.add_heading(s[level:].strip(), level=min(level, 3))
            i += 1

        elif s in HR:
            add_rule(doc)
            i += 1

        elif s.startswith(">"):
            # Join a wrapped blockquote into one paragraph.
            parts = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                parts.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_quote(doc, " ".join(p for p in parts if p), size)

        elif m := BULLET.match(line):
            indent, text = m.groups()
            i += 1
            # Hard-wrapped continuation lines belong to this bullet.
            while i < len(lines) and not starts_block(lines[i]):
                text += " " + lines[i].strip()
                i += 1
            add_bullet(doc, text, size, level=1 if len(indent) >= 2 else 0)

        elif m := NUMBERED.match(line):
            text = m.group(2)
            i += 1
            while i < len(lines) and not starts_block(lines[i]):
                text += " " + lines[i].strip()
                i += 1
            add_numbered(doc, text, size)

        elif not s:
            i += 1

        else:
            # Prose: join every hard-wrapped line into one real paragraph. Emitting
            # one paragraph per source line is what made these read as fragments.
            parts = [s]
            i += 1
            while i < len(lines) and not starts_block(lines[i]):
                parts.append(lines[i].strip())
                i += 1
            add_prose(doc, " ".join(parts), size)

    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


def main() -> None:
    args = sys.argv[1:]
    compact = "--compact" in args
    targets = [a for a in args if not a.startswith("--")] or DEFAULTS
    for name in targets:
        path = Path(name)
        if not path.is_absolute():
            path = ROOT / name
        if not path.exists():
            print(f"skipped {path.name} (not found)")
            continue
        out = convert(path, compact=compact)
        print(f"{path.name} -> {out.name} ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
